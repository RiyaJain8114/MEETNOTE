from fastapi import FastAPI, HTTPException, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
import whisper
import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import requests
import json

app = FastAPI()

class ChatRequest(BaseModel):
    path: str

# Add CORS middleware
origins = [
    "http://localhost:3000",
    "http://localhost:8000",
]

API_KEY = "AIzaSyCoR01KRzN_Fy2WV960UsN6sKX-kWG6y-o"
url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={API_KEY}"

# Request headers
headers = {
    "Content-Type": "application/json"
}


app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
def format_meeting_minutes(text):
    document = Document()
    document.add_heading('Minutes of Meeting', level=1)
    sections = re.split(r'# (.+):', text)  

    for i in range(len(sections)):
        if i % 2 == 1:
            heading = sections[i]
            paragraph = document.add_paragraph()
            run = paragraph.add_run(heading + ":")
            run.bold = True
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            paragraph = document.add_paragraph(sections[i].strip())
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    return document

temp_dir = "./temp"
os.makedirs(temp_dir, exist_ok=True)

@app.post("/process/")
async def process_audio(file: UploadFile = File(...)):
    try:
        file_path = os.path.join(temp_dir, file.filename)
        
        with open(file_path, "wb") as f:
            f.write(await file.read())

        print("File saved:", file_path)
        
        model = whisper.load_model("base")
        result = model.transcribe(file_path, language="en")

        response = requests.post(url, headers=headers, data=json.dumps({
            "contents": [
                {
                    "parts": [
                        {
                            "text": f"""
                            You are an expert executive assistant and note-taker.
                            I will provide a raw transcript of a meeting. Your job is to extract and format a clear, concise, and professional **Minutes of Meeting (MoM)** document.
                            
                            Please format the output exactly in this structure:
                            - Use `# Heading:` for each section (e.g., `# Meeting Title:`)
                            - Put the relevant content immediately below the heading
                            - Use bullet points where appropriate (like for attendees, agenda, discussion points)
                            - No markdown formatting like `**` or `*`, just plain text
                            - No blank lines between heading and content
                            - Each section must be on a new line
                            - Keep content formal and professional

                            Include these structured sections in your output:
                            1. **Meeting Title:** (Extract from context or summarize appropriately)
                            2. **Date and Time:** (Extract if available; else mention "Not specified")
                            3. **Attendees:** (List only actual names or roles mentioned in the transcript)
                            4. **Agenda:** (Summarize or infer from the transcript)
                            5. **Discussion Points:** (Summarize the main discussions clearly and concisely in bullet points)
                            6. **Decisions Made:** (Highlight any decisions taken during the meeting)
                            7. **Action Items:** (Use a table format with `Task`, `Owner`, and `Deadline` if mentioned)
                            8. **Next Steps / Follow-ups:** (Only include if discussed)

                            Make the language formal and business-appropriate, and avoid redundancy or vague placeholders. If any section has no information, omit it entirely rather than adding "not discussed".

                            Here is the transcript:
                            {result['text']}
                            """
                        }
                    ]
                }
            ]
        }))

        mom_result = response.json()["candidates"][0]["content"]["parts"][0]["text"]

        summary_response = requests.post(url, headers=headers, data=json.dumps({
            "contents": [
                {
                    "parts": [
                        {
                            "text": f"""
                            You are an expert assistant skilled at summarizing business meetings.
                            I will provide you with a raw transcript of a meeting. Your task is to generate a concise summary of the meeting.
                            Please include the key points discussed, decisions made, and any action items.

                            ### Transcript :
                            {result['text']}
                            """
                        }
                    ]
                }
            ]
        }))

        mom_result_summary = summary_response.json()["candidates"][0]["content"]["parts"][0]["text"]

        doc_mom = format_meeting_minutes(mom_result)
        mom_path = os.path.join(temp_dir, f"{os.path.splitext(file.filename)[0]}_Minutes_of_Meeting.docx")
        doc_mom.save(mom_path)

        doc_summary = Document()
        doc_summary.add_heading('Summary', level=1)
        doc_summary.add_paragraph(mom_result_summary)
        summary_path = os.path.join(temp_dir, f"{os.path.splitext(file.filename)[0]}_Summary.docx")
        doc_summary.save(summary_path)

        os.remove(file_path)
        
        return {"minutes_of_meeting_path": mom_path, "summary_path": summary_path}
    except Exception as e:
        print("Error processing audio:", e)
        raise HTTPException(status_code=500, detail=f"Error processing audio: {e}")

@app.get("/download")
async def download_file(path: str):
    if os.path.exists(path):
        return FileResponse(path, filename=os.path.basename(path))
    else:
        raise HTTPException(status_code=404, detail="File not found")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
